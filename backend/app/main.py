from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Query, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse, Response
from fastapi.staticfiles import StaticFiles
import pypdf
import docx
import os
import shutil
import uuid
import sqlite3
import datetime
import json
import re
from typing import Optional, List, Literal
from pydantic import BaseModel, Field
import httpx
print(f"DEBUG: main.py loaded from {os.path.abspath(__file__)}")

# 导入本地模块
UPLOAD_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)
FRONTEND_DIST_DIR = os.path.normpath(
    os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "..", "frontend", "dist")
)

from app.auth import auth_router
from app.database import (
    DB_PATH,
    init_db,
    migrate_db,
    get_dashboard_data, 
    get_conversations, 
    get_conversation_messages, 
    save_message,
    create_conversation,
    rename_conversation,
    delete_conversation
)
from langchain_core.messages import HumanMessage
from app.agent.graph import app_graph
from app.agent.reasoning_graph_store import (
    build_reasoning_graph_query,
    get_neo4j_connect_url,
    get_neo4j_browser_origin,
    store_reasoning_graph_in_neo4j,
)
from app.ingestion.db_config import db

app = FastAPI(
    title="VentureAgent API",
    description="VentureAgent MVP Backend API",
    version="0.1.0",
)

# 配置 CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.include_router(auth_router, prefix="/api/auth", tags=["auth"])

app.mount("/api/uploads", StaticFiles(directory=UPLOAD_DIR), name="uploads")

# --- Models ---

class ChatRequest(BaseModel):
    message: str
    session_id: Optional[str] = "default_session_123"
    project_id: Optional[int] = None
    agent: Optional[Literal["learning_tutor", "project_coach"]] = None

class ChatResponse(BaseModel):
    reply: str
    agent: str
    message_id: Optional[int] = None
    reasoning_trace: Optional[str] = None
    reasoning_graph: Optional[dict] = None


def _browser_proxy_base(request: Request) -> str:
    return str(request.base_url).rstrip("/") + "/api/neo4j-browser"


def _rewrite_browser_html(html: str, request: Request, message_id: int) -> str:
    browser_base = _browser_proxy_base(request)
    query = build_reasoning_graph_query(message_id)
    connect_url = get_neo4j_connect_url()
    bootstrap_config = f"""
<script>
window.__VA_BROWSER_BOOTSTRAP__ = {{
  initCmd: {json.dumps(query)},
  playImplicitInitCommands: true,
  connectURL: {json.dumps(connect_url)},
  db: "neo4j",
  cmd: "edit",
  arg: {json.dumps(query)}
}};
(function() {{
  const current = new URL(window.location.href);
  current.searchParams.set('cmd', 'edit');
  current.searchParams.set('arg', {json.dumps(query)});
  current.searchParams.set('db', 'neo4j');
  current.searchParams.set('connectURL', {json.dumps(connect_url)});
  window.history.replaceState(null, '', current.toString());
  try {{
    window.localStorage.setItem('neo4j.settings', JSON.stringify({{
      maxHistory: 30,
      theme: 'auto',
      initCmd: {json.dumps(query)},
      playImplicitInitCommands: true
    }}));
  }} catch (error) {{
    console.warn('Failed to seed neo4j.settings', error);
  }}
}})();
</script>
"""
    if "<head>" in html:
        html = html.replace("<head>", f'<head><base href="{browser_base}/assets/">{bootstrap_config}', 1)
    html = re.sub(
        r'((?:href|src)=["\'])((?!https?:|data:|#|/api/neo4j-browser)([^"\']+))(["\'])',
        lambda match: f"{match.group(1)}{browser_base}/assets/{match.group(3)}{match.group(4)}",
        html,
    )
    return html


async def _proxy_browser_request(
    request: Request,
    upstream_path: str,
    *,
    message_id: Optional[int] = None,
) -> Response:
    browser_origin = get_neo4j_browser_origin()
    target_url = f"{browser_origin.rstrip('/')}/{upstream_path.lstrip('/')}"
    if request.url.query:
        target_url = f"{target_url}?{request.url.query}"

    headers = {}
    for key, value in request.headers.items():
        lower_key = key.lower()
        if lower_key in {"host", "content-length"}:
            continue
        headers[key] = value

    async with httpx.AsyncClient(follow_redirects=True, timeout=30.0) as client:
        upstream_response = await client.request(
            request.method,
            target_url,
            headers=headers,
            content=await request.body(),
        )

    content = upstream_response.content
    media_type = upstream_response.headers.get("content-type", "text/plain")
    response_headers = {
        key: value
        for key, value in upstream_response.headers.items()
        if key.lower() not in {"content-length", "content-security-policy", "x-frame-options"}
    }
    response_headers["Cache-Control"] = "no-store"

    if "text/html" in media_type and message_id is not None:
        html = content.decode(upstream_response.encoding or "utf-8", errors="replace")
        html = _rewrite_browser_html(html, request, message_id)
        return HTMLResponse(content=html, status_code=upstream_response.status_code, headers=response_headers)

    return Response(
        content=content,
        status_code=upstream_response.status_code,
        headers=response_headers,
        media_type=media_type,
    )

class CreateConvRequest(BaseModel):
    id: str
    user_id: str
    title: str
    greeting: str

class RenameRequest(BaseModel):
    title: str

class MemberInfo(BaseModel):
    name: str
    student_id: Optional[str] = None
    role: str = "Member"
    position: Optional[str] = "队员"
    college: Optional[str] = None
    major: Optional[str] = None
    grade: Optional[str] = None
    info: Optional[str] = None

class ProjectCreateRequest(BaseModel):
    id: Optional[int] = None  # 用于区分更新操作
    name: str
    owner_id: str
    competition: Optional[str] = "互联网+"
    track: Optional[str] = None
    college: Optional[str] = None
    advisorName: Optional[str] = None
    advisorId: Optional[str] = None
    advisorInfo: Optional[str] = None
    content: Optional[str] = ""
    members: List[MemberInfo] = []

class StudentCapabilityScores(BaseModel):
    innovation: int = Field(description="创新性评分，0 到 100")
    feasibility: int = Field(description="落地性评分，0 到 100")
    technology: int = Field(description="技术力评分，0 到 100")
    team_fit: int = Field(description="团队契合评分，0 到 100")
    market: int = Field(description="市场潜力评分，0 到 100")
    compliance: int = Field(description="合规性评分，0 到 100")

class StudentCapabilityProfile(BaseModel):
    scores: StudentCapabilityScores
    comment: str = Field(description="针对短板的短评，必须引用学生原话")

# --- Helper Logic ---

def verify_project_members(project: ProjectCreateRequest):
    """通用校验逻辑：验证负责人、指导老师及团队成员在 Neo4j 中的实名状态"""
    # 1. 组长约束判定
    has_leader = any(m.role in ['Leader', '组长', '队长'] for m in project.members)
    if not has_leader:
        raise HTTPException(status_code=400, detail="项目必须包含至少一名组长（Leader/组长）")

    # 2. 指导老师实名核验 (Neo4j)
    if project.advisorName:
        verify_advisor_q = """
            MATCH (u:User {role: 'teacher', real_name: $name, username: $id, college: $college}) 
            RETURN u
        """
        advisor_res = db.execute_query(verify_advisor_q, {
            "name": project.advisorName, 
            "id": project.advisorId, 
            "college": project.college
        })
        if not advisor_res:
            raise HTTPException(status_code=400, detail=f"指导老师校验失败：无法在系统中找到名为 {project.advisorName}、工号为 {project.advisorId} 且属于 {project.college} 的已注册教师。")

    # 3. 团队成员实名核验 (Neo4j)
    for m in project.members:
        verify_member_q = """
            MATCH (u:User {role: 'student', real_name: $name, username: $id, college: $college, major: $major, grade: $grade}) 
            RETURN u
        """
        member_res = db.execute_query(verify_member_q, {
            "name": m.name,
            "id": m.student_id,
            "college": m.college or project.college,
            "major": m.major,
            "grade": m.grade
        })
        if not member_res:
            raise HTTPException(status_code=400, detail=f"成员校验失败：名为 {m.name}、学号为 {m.student_id} 的同学未注册或填写的学院专业年级信息有误。")

class CommitRequest(BaseModel):
    content: str
    author: Optional[str] = "System"

class ReviewRequest(BaseModel):
    rubric: str = "internet_plus"

class UpdateContentRequest(BaseModel):
    content: str
    file_id: Optional[int] = None

class TeacherInterventionRequest(BaseModel):
    project_id: int
    teacher_name: str
    content: str

class AuditResult(BaseModel):
    r1: float = Field(description="R1 问题定义评分，1 到 5")
    r2: float = Field(description="R2 用户证据强度评分，1 到 5")
    r3: float = Field(description="R3 方案可行性评分，1 到 5")
    r4: float = Field(description="R4 商业模式一致性评分，1 到 5")
    r5: float = Field(description="R5 市场与竞争评分，1 到 5")
    r6: float = Field(description="R6 财务健康逻辑评分，1 到 5")
    r7: float = Field(description="R7 创新与差异化评分，1 到 5")
    r8: float = Field(description="R8 团队与执行力评分，1 到 5")
    r9: float = Field(description="R9 表达与材料质量评分，1 到 5")
    overall_risk: Literal["High", "Medium", "Low"] = Field(description="总体风险评级")
    audit_summary: str = Field(description="给指导老师看的深度审计结论，150字以内")
    evidence_trace: str = Field(description="必须引用项目原文中的1到3处关键片段作为证据，说明这些原文为什么支撑上述判断。")

def extract_h_principles(*texts: str) -> list[str]:
    """
    从审计文本中提取 H 原则标签。
    优先识别显式的 H1-H9 标签；若历史数据未输出标签，则使用关键词兜底推断。
    """
    import re

    combined_text = "\n".join(text for text in texts if text)
    if not combined_text.strip():
        return []

    explicit_tags = []
    seen = set()
    for tag in re.findall(r"(H[1-9])", combined_text):
        if tag not in seen:
            explicit_tags.append(tag)
            seen.add(tag)
    if explicit_tags:
        return explicit_tags

    keyword_map = {
        "H1": ["价值主张错位", "定位混乱", "内容脱节", "标题与内容", "项目定位", "概念混淆", "套壳", "错位"],
        "H2": ["用户细分错误", "目标用户过宽", "用户画像不清", "客户群体过宽", "用户细分"],
        "H3": ["痛点假设", "伪需求", "痛点不明确", "需求不强", "需求存疑"],
        "H4": ["市场规模", "百亿市场", "市场空间", "规模推断", "市场过大"],
        "H5": ["缺乏验证", "用户调研", "市场验证", "小规模测试", "验证数据", "可行性验证", "反馈不足"],
        "H6": ["竞争壁垒", "壁垒薄弱", "仅依赖专利", "缺乏壁垒", "护城河", "商业化验证不足"],
        "H7": ["推广渠道", "渠道错位", "获客渠道", "渠道策略", "推广路径"],
        "H8": ["单位经济", "收入模型", "财务预测", "获客成本", "毛利", "盈利模型", "LTV", "CAC", "计算依据", "营收逻辑", "成本结构"],
        "H9": ["团队资源", "资源错配", "执行力不足", "团队能力", "核心成员缺失"],
    }

    inferred_tags = []
    for tag, keywords in keyword_map.items():
        if any(keyword in combined_text for keyword in keywords):
            inferred_tags.append(tag)
    return inferred_tags

def build_teacher_top_mistakes(cursor, teacher_identity: str):
    """
    汇总教师名下项目的 H 原则高频错误，供教师大盘与教学计划生成共用。
    teacher_identity 同时兼容教师工号与姓名。
    """
    cursor.execute("""
        SELECT pa.audit_summary, pa.evidence_trace
        FROM project_assessments pa
        JOIN projects p ON pa.project_id = p.id
        WHERE (p.advisor_id = ? OR p.advisor_name = ?)
          AND pa.audit_summary IS NOT NULL
    """, (teacher_identity, teacher_identity))
    assessment_rows = cursor.fetchall()

    from collections import Counter

    h_counts = Counter()
    for row in assessment_rows:
        found_tags = set(extract_h_principles(row["audit_summary"], row["evidence_trace"]))
        for tag in found_tags:
            h_counts[tag] += 1

    h_desc_map = {
        "H1": "价值主张错位", "H2": "用户细分错误", "H3": "痛点假设不成立",
        "H4": "市场规模幻觉", "H5": "缺乏可行性验证", "H6": "竞争壁垒薄弱",
        "H7": "推广渠道错位", "H8": "单位经济不成立", "H9": "团队资源错配"
    }

    top_mistakes = []
    for tag, count in h_counts.most_common(5):
        top_mistakes.append({
            "tag": tag,
            "name": f"{tag} {h_desc_map.get(tag, '未定义原则')}",
            "count": count,
            "desc": f"在{count}个项目中出现该原则报错"
        })

    return top_mistakes

def resolve_project_audit_context(project_id: int):
    """
    解析教师审计与学生手动评审所依赖的项目正文。
    优先使用最近一次上传文件的 content；若不存在，则回退到 projects.content。
    """
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    try:
        cursor.execute("SELECT id, name, content, competition FROM projects WHERE id = ?", (project_id,))
        project = cursor.fetchone()
        if not project:
            return None

        cursor.execute("""
            SELECT id, filename, content, created_at
            FROM project_files
            WHERE project_id = ?
            ORDER BY datetime(created_at) DESC, id DESC
            LIMIT 1
        """, (project_id,))
        latest_file = cursor.fetchone()

        latest_file_content = ""
        latest_file_name = None
        if latest_file:
            latest_file_content = (latest_file["content"] or "").strip()
            latest_file_name = latest_file["filename"]

        project_content = (project["content"] or "").strip()
        resolved_content = latest_file_content or project_content

        return {
            "project_id": project["id"],
            "name": project["name"],
            "competition": project["competition"],
            "content": resolved_content,
            "content_source": latest_file_name or "项目正文",
        }
    finally:
        conn.close()

def clamp_percentage_score(value) -> int:
    try:
        numeric = int(round(float(value)))
    except (TypeError, ValueError):
        numeric = 0
    return max(0, min(100, numeric))

def extract_json_payload(raw_content) -> dict:
    text = raw_content
    if isinstance(text, list):
        text = "".join(str(part) for part in text)
    text = str(text).strip()

    if "```json" in text:
        text = text.split("```json", 1)[1].split("```", 1)[0].strip()
    elif "```" in text:
        text = text.split("```", 1)[1].split("```", 1)[0].strip()

    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        text = text[start:end + 1]

    payload = json.loads(text)
    if not isinstance(payload, dict):
        raise ValueError("LLM 未返回 JSON 对象")
    return payload

def normalize_evidence_trace(text: str) -> str:
    normalized = str(text or "").replace("\r", "")
    normalized = normalized.replace("<br/>", "\n").replace("<br />", "\n").replace("<br>", "\n")
    normalized = normalized.strip()
    if not normalized:
        return ""

    import re

    normalized = re.sub(r"\s+(?=(?:\d+[\.、\)])|(?:[①②③④⑤⑥⑦⑧⑨⑩]))", "\n\n", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()

# --- Endpoints ---

@app.on_event("startup")
def on_startup():
    init_db()
    migrate_db()
    
    # # 注入 Mock 教师用户到账户库 (同步更新 username 与 real_name)
    # from app.ingestion.db_config import db
    # try:
    #     # T001 对应 张老师
    #     db.execute_query("MERGE (u:User {username: 'T001'}) SET u.password = '123456', u.role = 'teacher', u.real_name = '张老师', u.college = '智慧双创学院'")
    #     # T002 对应 王老师
    #     db.execute_query("MERGE (u:User {username: 'T002'}) SET u.password = '123456', u.role = 'teacher', u.real_name = '王老师', u.college = '经管学院'")
        
    #     # 为了兼容用户可能的输入习惯，也确保 '张老师' 这个 username 有 real_name 自指
    #     db.execute_query("MERGE (u:User {username: '张老师'}) SET u.password = '123456', u.role = 'teacher', u.real_name = '张老师', u.college = '智慧双创学院'")

    #     # 注入 Mock 管理员用户
    #     db.execute_query("MERGE (u:User {username: 'admin'}) SET u.password = 'admin123', u.role = 'admin', u.real_name = '超级管理员', u.college = '校级管理中心'")
    #     db.execute_query("MERGE (u:User {username: '001'}) SET u.password = '123456', u.role = 'admin', u.real_name = '管理员001', u.college = '系统维护部'")
    # except Exception as e:
    #     print(f"Mock Teacher Injection failed (Neo4j possibly down): {e}")

    # # --- SQLite 数据对齐逻辑：将 advisor_name 映射回已知的 advisor_id ---
    # try:
    #     conn = sqlite3.connect(DB_PATH)
    #     cursor = conn.cursor()
    #     # 兼容性回填：同时支持 T001 和直接使用姓名作为 ID 的情况
    #     cursor.execute("UPDATE projects SET advisor_id = 'T001' WHERE advisor_name = '张老师'")
    #     cursor.execute("UPDATE projects SET advisor_id = '张老师' WHERE advisor_name = '张老师'") 
    #     cursor.execute("UPDATE projects SET advisor_id = 'T002' WHERE advisor_name = '王老师'")
    #     cursor.execute("UPDATE projects SET advisor_id = '王老师' WHERE advisor_name = '王老师'")
    #     conn.commit()
    #     conn.close()
    #     print("DEBUG: SQLite advisor_id comprehensive backfill completed.")
    # except Exception as e:
    #     print(f"SQLite backfill failed: {e}")

@app.get("/")
def read_root():
    index_path = os.path.join(FRONTEND_DIST_DIR, "index.html")
    if os.path.isfile(index_path):
        return FileResponse(index_path)
    return {"message": "Welcome to VentureAgent API"}

@app.get("/health")
def health_check():
    return {"status": "ok"}

@app.get("/api/sync/dashboard")
def sync_dashboard(user_id: str):
    data = get_dashboard_data(user_id)
    return data if data else {"projects": [], "deadlines": [], "evolution_logs": [], "members": []}

@app.get("/api/conversations")
def list_conversations(user_id: str):
    return get_conversations(user_id)

@app.get("/api/conversations/{conv_id}/messages")
def list_messages(conv_id: str):
    return get_conversation_messages(conv_id)


@app.get("/api/neo4j-browser/view/{message_id}")
async def neo4j_browser_view(message_id: int, request: Request):
    return await _proxy_browser_request(request, "/browser/", message_id=message_id)


@app.api_route("/api/neo4j-browser/assets/{asset_path:path}", methods=["GET", "HEAD"])
async def neo4j_browser_assets(asset_path: str, request: Request):
    return await _proxy_browser_request(request, f"/browser/{asset_path}")


@app.api_route("/api/neo4j-browser/{asset_name}", methods=["GET", "HEAD"])
async def neo4j_browser_asset_fallback(asset_name: str, request: Request):
    if "/" in asset_name:
        raise HTTPException(status_code=404, detail="Not found")
    if not re.search(r"\.(js|css|json|png|jpg|jpeg|svg|ico|map|woff2?|ttf)$", asset_name, re.IGNORECASE):
        raise HTTPException(status_code=404, detail="Not found")
    return await _proxy_browser_request(request, f"/browser/{asset_name}")


@app.api_route("/api/neo4j-browser/db/{db_path:path}", methods=["GET", "POST", "PUT", "PATCH", "DELETE", "OPTIONS", "HEAD"])
async def neo4j_browser_db_proxy(db_path: str, request: Request):
    return await _proxy_browser_request(request, f"/db/{db_path}")

@app.post("/api/conversations")
def create_conv(request: CreateConvRequest):
    create_conversation(request.id, request.user_id, request.title, request.greeting)
    return {"status": "ok"}

@app.patch("/api/conversations/{conv_id}")
def rename_conv(conv_id: str, request: RenameRequest):
    rename_conversation(conv_id, request.title)
    return {"status": "ok"}

@app.delete("/api/conversations/{conv_id}")
def delete_conv(conv_id: str):
    try:
        delete_conversation(conv_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"删除会话失败: {str(e)}")
    return {"status": "ok"}

@app.get("/api/conversations/{conv_id}/capability-profile")
async def generate_conversation_capability_profile(conv_id: str):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    try:
        cursor.execute("SELECT id, title FROM conversations WHERE id = ?", (conv_id,))
        conversation = cursor.fetchone()
        if not conversation:
            raise HTTPException(status_code=404, detail="会话不存在")

        cursor.execute("""
            SELECT id, role, agent, content, timestamp
            FROM messages
            WHERE conversation_id = ?
              AND NOT (role = 'coach' AND COALESCE(agent, '') = '系统助手')
            ORDER BY datetime(timestamp) DESC, id DESC
            LIMIT 6
        """, (conv_id,))
        recent_messages = cursor.fetchall()
    finally:
        conn.close()

    if len(recent_messages) < 6:
        raise HTTPException(status_code=400, detail="至少需要 3 轮问答才能生成能力画像")

    recent_messages = list(reversed(recent_messages))
    transcript_lines = []
    for index, row in enumerate(recent_messages):
        round_index = index // 2 + 1
        speaker = "学生" if row["role"] == "user" else (row["agent"] or "AI导师")
        transcript_lines.append(f"第{round_index}轮{speaker}：{row['content']}")
    transcript = "\n".join(transcript_lines)

    try:
        from app.agent.graph import get_llm

        llm = get_llm()
        prompt = f"""你是一名高校创新创业课程导师。下面是某位学生在当前会话中最近 3 轮问答（共 6 条消息）的原始记录，请只基于这些内容完成分析，不要使用会话外信息。

会话标题：{conversation["title"]}
对话记录：
{transcript}

请严格输出 JSON，对象格式如下：
{{
  "scores": {{
    "innovation": 0到100的整数,
    "feasibility": 0到100的整数,
    "technology": 0到100的整数,
    "team_fit": 0到100的整数,
    "market": 0到100的整数,
    "compliance": 0到100的整数
  }},
  "comment": "80到160字的中文短评"
}}

评分维度含义如下：
- innovation：创新性
- feasibility：落地性
- technology：技术力
- team_fit：团队契合
- market：市场潜力
- compliance：合规性

补充要求：
1. comment 只聚焦最明显的 1 到 2 个短板。
2. comment 必须引用至少一句学生原话作为证据，引用时使用中文引号『』。
3. comment 必须明确指出“第一轮 / 第二轮 / 第三轮”中的至少一轮，例如“学生在第一轮提到：『我们的客户是所有人』”。
4. comment 要给出一条可以立刻执行的改进建议。
5. 只输出 JSON，不要输出 Markdown，不要补充解释。
"""
        response = llm.invoke(prompt)
        result = extract_json_payload(response.content)
        raw_scores = result.get("scores", {})
        normalized_scores = StudentCapabilityScores(
            innovation=clamp_percentage_score(raw_scores.get("innovation")),
            feasibility=clamp_percentage_score(raw_scores.get("feasibility")),
            technology=clamp_percentage_score(raw_scores.get("technology")),
            team_fit=clamp_percentage_score(raw_scores.get("team_fit")),
            market=clamp_percentage_score(raw_scores.get("market")),
            compliance=clamp_percentage_score(raw_scores.get("compliance")),
        )
        comment = str(result.get("comment", "")).strip()
        if not comment:
            raise ValueError("画像短评为空")

        profile = StudentCapabilityProfile(scores=normalized_scores, comment=comment)
        return {"status": "ok", "profile": profile.model_dump()}
    except HTTPException:
        raise
    except Exception as e:
        print(f"Failed to generate conversation capability profile: {e}")
        raise HTTPException(status_code=500, detail=f"能力画像生成失败：{str(e)}")

@app.post("/api/chat", response_model=ChatResponse)
def chat_endpoint(request: ChatRequest):
    # --- 教师干预逻辑注入 ---
    final_input_content = request.message
    if request.project_id:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        cursor.execute("SELECT content FROM teacher_interventions WHERE project_id = ? AND is_active = 1", (request.project_id,))
        interventions = cursor.fetchall()
        conn.close()
        if interventions:
            rules_text = " | ".join([row['content'] for row in interventions])
            # 将干预指令作为系统提示注入 (此处简单包裹在 HumanMessage 前面)
            final_input_content = f"【教师干预约束：{rules_text}】\n用户的实际消息：{request.message}"

    input_message = HumanMessage(content=final_input_content)
    initial_state = {
        "messages": [input_message],
        "next_agent": "",
        "forced_agent": request.agent or ""
    }
    session_id = request.session_id or "default_session_123"
    config = {"configurable": {"thread_id": session_id}}
    
    save_message(session_id, "user", request.message)
    final_state = app_graph.invoke(initial_state, config=config)
    
    messages = final_state.get("messages", [])
    final_reply = messages[-1].content if messages else "抱歉，系统内部状态流转异常，未返回消息。"
    context = final_state.get("context", {}) or {}
    reasoning_trace = context.get("reasoning_trace")
    reasoning_graph = context.get("reasoning_graph")
    
    agent_id = final_state.get("next_agent", "unknown")
    agent_map = {
        "learning_tutor": "学习辅导 Agent (A1)",
        "project_coach": "项目教练 Agent (A2)",
        "unknown": "系统导师"
    }
    display_name = agent_map.get(agent_id, "系统导师")
    
    if isinstance(final_reply, list):
        final_reply = "".join(str(part) for part in final_reply)
        
    coach_message_id = save_message(
        session_id,
        "coach",
        str(final_reply),
        display_name,
        reasoning_trace=str(reasoning_trace) if reasoning_trace else None,
        reasoning_graph=reasoning_graph if isinstance(reasoning_graph, dict) else None
    )
    if coach_message_id and isinstance(reasoning_graph, dict):
        try:
            store_reasoning_graph_in_neo4j(
                message_id=int(coach_message_id),
                conversation_id=session_id,
                agent=display_name,
                reasoning_graph=reasoning_graph,
            )
        except Exception as exc:
            print(f"Failed to sync reasoning_graph to Neo4j: {exc}")

    return ChatResponse(
        reply=str(final_reply),
        agent=display_name,
        message_id=int(coach_message_id) if coach_message_id else None,
        reasoning_trace=str(reasoning_trace) if reasoning_trace else None,
        reasoning_graph=reasoning_graph if isinstance(reasoning_graph, dict) else None
    )

@app.post("/api/projects")
async def create_project_endpoint(project: ProjectCreateRequest):
    # 严格实名准入核验 (通用逻辑已覆盖创建与更新)
    verify_project_members(project)

    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # 判断是创建还是更新逻辑
    if project.id:
        # 更新存量项目基本信息
        cursor.execute("""
            UPDATE projects SET name=?, competition=?, track=?, college=?, advisor_name=?, advisor_id=?, advisor_info=?
            WHERE id=?
        """, (project.name, project.competition, project.track, project.college, project.advisorName, project.advisorId, project.advisorInfo, project.id))
        project_id = project.id
        # 清理并重建成员关系（确保实名绑定的最新性）
        cursor.execute("DELETE FROM members WHERE project_id = ?", (project_id,))
    else:
        # 创建新项目
        cursor.execute("""
            INSERT INTO projects (name, owner_id, competition, track, college, advisor_name, advisor_id, advisor_info, content)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (project.name, project.owner_id, project.competition, project.track, project.college, project.advisorName, project.advisorId, project.advisorInfo, project.content))
        project_id = cursor.lastrowid
        
        today = datetime.date.today()
        ddls = []
        if project.competition == "互联网+":
            ddls = [
                ("项目立项", today.strftime("%Y-%m-%d")),
                ("计划书初稿截止", (today + datetime.timedelta(days=14)).strftime("%Y-%m-%d")),
                ("商业模式深度复核", (today + datetime.timedelta(days=28)).strftime("%Y-%m-%d")),
                ("国赛/省赛网评截止", "2026-06-30")
            ]
        elif project.competition == "挑战杯":
            ddls = [
                ("技术方案定稿", (today + datetime.timedelta(days=10)).strftime("%Y-%m-%d")),
                ("实物模型演示", (today + datetime.timedelta(days=30)).strftime("%Y-%m-%d")),
                ("省赛申报截止", "2026-04-15")
            ]
        else:
            ddls = [
                ("项目启动", today.strftime("%Y-%m-%d")),
                ("中期进度汇报", (today + datetime.timedelta(days=30)).strftime("%Y-%m-%d")),
                ("结项验收", (today + datetime.timedelta(days=60)).strftime("%Y-%m-%d"))
            ]
        for title, d_date in ddls:
            cursor.execute("INSERT INTO deadlines (project_id, title, due_date) VALUES (?, ?, ?)", (project_id, title, d_date))
        
        cursor.execute("INSERT INTO evolution_logs (project_id, event) VALUES (?, ?)", (project_id, f"完成《{project.name}》初步申报入库"))
    
    # 插入项目成员表
    if project.advisorName:
        cursor.execute("""
            INSERT INTO members (project_id, name, student_id, role, position, info, college)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (project_id, project.advisorName, project.advisorId, "Advisor", "指导老师", project.advisorInfo, project.college))

    for m in project.members:
        cursor.execute("""
            INSERT INTO members (project_id, name, student_id, college, major, grade, role, position, info)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (project_id, m.name, m.student_id, m.college or project.college, m.major, m.grade, m.role, m.position, m.info))

    conn.commit()
    conn.close()
    return {"status": "ok", "project_id": project_id}

@app.post("/api/projects/import")
async def import_project_file(file: UploadFile = File(...), project_id: Optional[int] = Form(None)):
    content = ""
    filename = file.filename or "unknown_file"
    unique_filename = f"{uuid.uuid4().hex}_{filename}"
    upload_path = os.path.join(UPLOAD_DIR, unique_filename)
    
    with open(upload_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
        
    file.file.seek(0)
    file_url = f"/api/uploads/{unique_filename}"
    try:
        if filename.endswith(".pdf"):
            pdf_reader = pypdf.PdfReader(file.file)
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text: content += text + "\n"
        elif filename.endswith(".docx"):
            from io import BytesIO
            doc = docx.Document(BytesIO(await file.read()))
            for para in doc.paragraphs: content += para.text + "\n"
        else:
            return {"error": "Unsupported file format. Please upload PDF or DOCX."}
    except Exception as e:
        return {"error": f"Failed to parse file: {str(e)}"}
    
    # 强制清理：移除文档中解析出来的非法/游离 surrogate characters (如 \ud835 等) 防止 JSON 序列化崩溃
    import re
    content = re.sub(r'[\ud800-\udfff]', '', content)
    
    if project_id is not None:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        cur = conn.cursor()
        cur.execute(
            """
            SELECT id, file_url
            FROM project_files
            WHERE project_id = ? AND filename = ?
            ORDER BY datetime(created_at) DESC, id DESC
            LIMIT 1
            """,
            (project_id, filename)
        )
        existing_file = cur.fetchone()

        if existing_file:
            old_file_url = existing_file["file_url"]
            try:
                filename_part = old_file_url.split("/api/uploads/")[-1]
                old_path = os.path.join(UPLOAD_DIR, filename_part)
                if os.path.exists(old_path):
                    os.remove(old_path)
            except Exception as cleanup_error:
                print(f"Failed to remove old uploaded file for overwrite: {cleanup_error}")

            cur.execute(
                """
                UPDATE project_files
                SET file_url = ?, content = ?, created_at = CURRENT_TIMESTAMP
                WHERE id = ? AND project_id = ?
                """,
                (file_url, content, existing_file["id"], project_id)
            )
            file_id = existing_file["id"]
        else:
            cur.execute(
                "INSERT INTO project_files (project_id, filename, file_url, content) VALUES (?, ?, ?, ?)",
                (project_id, filename, file_url, content)
            )
            file_id = cur.lastrowid
        conn.commit()
        conn.close()

        return {"status": "ok", "text": content, "filename": filename, "file_id": file_id, "file_url": f"/api/uploads/{unique_filename}"}

@app.get("/api/projects/{project_id}/files")
def get_project_files(project_id: int):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("""
        SELECT id, filename, file_url, created_at
        FROM project_files
        WHERE project_id = ?
        ORDER BY datetime(created_at) DESC, id DESC
    """, (project_id,))
    rows = cursor.fetchall()
    conn.close()
    return [dict(row) for row in rows]

@app.get("/api/projects/{project_id}/files/{file_id}")
async def get_project_file_content(project_id: int, file_id: int):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT id, filename, file_url, content FROM project_files WHERE id = ? AND project_id = ?", (file_id, project_id))
    row = cursor.fetchone()
    if not row:
        conn.close()
        return {"error": "File not found"}

    stored_content = (row["content"] or "").strip()
    if stored_content:
        conn.close()
        return {"id": file_id, "filename": row["filename"], "content": row["content"]}
    
    file_url = row["file_url"]
    filename = row["filename"]
    content = ""
    try:
        filename_part = file_url.split("/api/uploads/")[-1]
        physical_path = os.path.join(UPLOAD_DIR, filename_part)
        if filename.endswith(".pdf"):
            with open(physical_path, "rb") as f:
                pdf_reader = pypdf.PdfReader(f)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text: content += text + "\n"
        elif filename.endswith(".docx"):
            doc = docx.Document(physical_path)
            for para in doc.paragraphs: content += para.text + "\n"
        else:
            content = "Non-previewable file type."
    except Exception as e:
        content = f"Error reading file: {str(e)}"
    else:
        # 为历史文件做惰性回填，后续审计与预览统一走数据库中的 content
        cursor.execute("UPDATE project_files SET content = ? WHERE id = ? AND project_id = ?", (content, file_id, project_id))
        conn.commit()
    conn.close()
    return {"id": file_id, "filename": filename, "content": content}

@app.delete("/api/project-files/{file_id}")
def delete_project_file(file_id: int):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM project_files WHERE id = ?", (file_id,))
    row = cursor.fetchone()
    if not row:
        conn.close()
        return {"error": "File not found"}
    file_url = row["file_url"]
    try:
        filename_part = file_url.split("/api/uploads/")[-1]
        physical_path = os.path.join(UPLOAD_DIR, filename_part)
        if os.path.exists(physical_path): os.remove(physical_path)
    except Exception: pass
    cursor.execute("DELETE FROM project_files WHERE id = ?", (file_id,))
    conn.commit()
    conn.close()
    return {"status": "deleted"}

@app.post("/api/projects/{project_id}/commits")
def create_commit(project_id: int, request: CommitRequest):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("INSERT INTO project_commits (project_id, content, author) VALUES (?, ?, ?)", 
                   (project_id, request.content, request.author))
    conn.commit()
    conn.close()
    return {"status": "ok"}

@app.post("/api/projects/{project_id}/review")
def review_project(project_id: int, request: ReviewRequest):
    audit_context = resolve_project_audit_context(project_id)
    if not audit_context:
        return {"review": "项目未找到或正文为空。"}

    project_name = audit_context["name"]
    project_content = audit_context["content"]
    if not project_content or len(project_content.strip()) < 10:
        return {"review": "项目正文内容太少，AI 无法进行深度评估。请先在编辑器中补充内容。"}

    try:
        from app.agent.graph import get_llm
        llm = get_llm()
        rubric_name = "互联网+ (中国国际大学生创新大赛)" if request.rubric == "internet_plus" else "挑战杯 (全国大学生课外学术科技作品竞赛)"
        system_prompt = f"""你是一个资深的创业项目评审专家，正在为参加“{rubric_name}”的学生提供深度点评。
        请针对以下项目计划书正文进行多维度的客观评价..."""
        response = llm.invoke([("system", system_prompt), ("human", f"项目：{project_name}\n内容：{project_content}")])
        return {"review": response.content}
    except Exception as e:
        return {"review": f"AI 评审服务暂时不可用 ({str(e)})。"}

@app.put("/api/projects/{project_id}/content")
def update_project_content(project_id: int, request: UpdateContentRequest):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    if request.file_id:
        cursor.execute(
            "UPDATE project_files SET content = ? WHERE id = ? AND project_id = ?",
            (request.content, request.file_id, project_id)
        )
    else:
        cursor.execute("UPDATE projects SET content = ? WHERE id = ?", (request.content, project_id))
    conn.commit()
    conn.close()
    return {"status": "ok"}

@app.patch("/api/projects/{project_id}")
def rename_project(project_id: int, request: RenameRequest):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("UPDATE projects SET name = ? WHERE id = ?", (request.title, project_id))
    conn.commit()
    conn.close()
    return {"status": "ok"}

@app.delete("/api/projects/{project_id}")
def delete_project_endpoint(project_id: int):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    # 1. 物理回收附件文件
    try:
        cursor.execute("SELECT file_url FROM project_files WHERE project_id = ?", (project_id,))
        files = cursor.fetchall()
        for f in files:
            file_url = f["file_url"]
            try:
                filename_part = file_url.split("/api/uploads/")[-1]
                physical_path = os.path.join(UPLOAD_DIR, filename_part)
                if os.path.exists(physical_path):
                    os.remove(physical_path)
            except Exception as fe:
                print(f"Failed to remove file {file_url}: {fe}")
    except Exception as e:
        print(f"Error querying project files for deletion: {e}")

    # 2. 清理数据库关联记录
    related_tables = [
        "project_files", 
        "project_assessments", 
        "teacher_interventions", 
        "members", 
        "deadlines", 
        "evolution_logs", 
        "project_commits"
    ]
    for table in related_tables:
        try:
            cursor.execute(f"DELETE FROM {table} WHERE project_id = ?", (project_id,))
        except Exception as te:
            print(f"Failed to clean table {table} for project {project_id}: {te}")

    # 3. 删除项目本体
    cursor.execute("DELETE FROM projects WHERE id = ?", (project_id,))
    
    conn.commit()
    conn.close()
    return {"status": "ok"}

# --- 教师端 API 接口 ---

@app.get("/api/teacher/dashboard")
async def get_teacher_dashboard(teacher_id: str):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    # 获取名下所有项目及基础评估 (按 advisor_id 过滤)
    cursor.execute("""
        SELECT p.id, p.name as name, p.owner_id as leader, p.college, p.competition,
               pa.overall_risk as risk_level, pa.audit_summary
        FROM projects p
        LEFT JOIN project_assessments pa ON p.id = pa.project_id
        WHERE p.advisor_id = ?
    """, (teacher_id,))
    projects_rows = cursor.fetchall()
    projects = [dict(row) for row in projects_rows]
    
    # 统计辅导学生总数
    cursor.execute("""
        SELECT COUNT(DISTINCT m.student_id) as student_count
        FROM members m
        JOIN projects p ON m.project_id = p.id
        WHERE p.advisor_id = ? AND m.role != 'Advisor'
    """, (teacher_id,))
    student_count = cursor.fetchone()['student_count'] or 0
    
    # 统计高频错误 (基于 H 原则)
    top_mistakes = build_teacher_top_mistakes(cursor, teacher_id)
    if not top_mistakes:
        top_mistakes = [
            {"name": "暂无高频错误", "count": 0, "desc": "当前暂无风险预警项目"}
        ]
    
    conn.close()
    return {
        "projects": projects, 
        "stats": {
            "student_count": student_count,
            "project_count": len(projects),
            "high_risk_count": len([p for p in projects if p.get('risk_level') == 'High'])
        },
        "top_mistakes": top_mistakes
    }

@app.post("/api/teacher/projects/{project_id}/audit")
async def trigger_ai_audit(project_id: int):
    """
    触发 DeepSeek 深度审计：读取项目内容，生成 R1-R9 评分及 H 原则审计
    """
    proj = resolve_project_audit_context(project_id)
    if not proj:
        raise HTTPException(status_code=404, detail="项目不存在")

    if not proj["content"] or len(proj["content"].strip()) < 10:
        raise HTTPException(status_code=400, detail="项目正文内容太少，无法执行深度审计")

    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    try:
        from app.agent.graph import get_llm
        llm = get_llm()
        structured_llm = llm.with_structured_output(AuditResult)

        prompt = f"""
你是一名创业大赛资深评委，同时也是指导教师的“教师顾问”。
请对下面的学生项目材料做一次严格、具体、可追溯的深度审计。

项目名称：{proj['name']}
项目类型：{proj['competition']}
项目正文：
{proj['content'][:8000]}

输出要求：
1. 对 R1-R9 分别给出 1 到 5 分，可使用小数。
2. overall_risk 只能是 High、Medium、Low 之一。
3. audit_summary 面向指导教师，150 字以内，重点指出项目在核心逻辑、商业模式、竞争壁垒、单位经济性等方面最关键的问题，并给出一条最值得优先干预的建议。
4. audit_summary 中必须显式写出 1 到 3 个最关键的 H 原则编号，例如 H1、H6、H8，建议放在开头或句中明确点出，便于后续统计高频商业逻辑盲区。
5. evidence_trace 必须填写，且不能是空字符串。
6. evidence_trace 必须直接引用项目原文中的 1 到 3 处关键片段，优先引用更完整、更长的原文，不要只摘一个短词；每条引用尽量达到 20 到 60 个汉字。
7. evidence_trace 必须严格使用编号分条书写，例如：
   1. 原文引用：……
      评审说明：……
   2. 原文引用：……
      评审说明：……
8. 每一条中的“评审说明”都要写得更充分，至少 40 到 80 个汉字，解释这段原文为什么支撑你的判断，而不是只写一句结论。
9. 如果材料不足，就明确指出“材料缺失”本身也是证据，但仍要按编号分条输出。
10. 不要编造项目中没有出现过的事实。
11. 只返回结构化结果，不要输出额外说明。
"""

        audit_result = structured_llm.invoke(prompt)
        if isinstance(audit_result, AuditResult):
            res_data = audit_result.model_dump()
        elif isinstance(audit_result, dict):
            res_data = audit_result
        else:
            raise ValueError("AI 审计结果格式异常")

        evidence_trace = normalize_evidence_trace(str(res_data.get("evidence_trace", "")).strip())
        if not evidence_trace:
            raise ValueError("AI 审计结果缺少 evidence_trace")
        res_data["evidence_trace"] = evidence_trace
        
        # 保存到数据库
        cursor.execute("""
            INSERT INTO project_assessments (project_id, r1_score, r2_score, r3_score, r4_score, r5_score, r6_score, r7_score, r8_score, r9_score, overall_risk, audit_summary, evidence_trace)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ON CONFLICT(project_id) DO UPDATE SET
                r1_score=excluded.r1_score, r2_score=excluded.r2_score, 
                r3_score=excluded.r3_score, r4_score=excluded.r4_score,
                r5_score=excluded.r5_score, r6_score=excluded.r6_score,
                r7_score=excluded.r7_score, r8_score=excluded.r8_score,
                r9_score=excluded.r9_score, overall_risk=excluded.overall_risk,
                audit_summary=excluded.audit_summary, evidence_trace=excluded.evidence_trace
        """, (
            project_id, res_data.get('r1', 0), res_data.get('r2', 0), 
            res_data.get('r3', 0), res_data.get('r4', 0), res_data.get('r5', 0),
            res_data.get('r6', 0), res_data.get('r7', 0),
            res_data.get('r8', 0), res_data.get('r9', 0),
            res_data.get('overall_risk', 'Medium'), res_data.get('audit_summary', ''),
            res_data.get('evidence_trace', '')
        ))
        conn.commit()
    except Exception as e:
        print(f"Audit failed: {e}")
        raise HTTPException(status_code=500, detail=f"AI 审计失败: {str(e)}")
    finally:
        conn.close()
    
    return {"status": "ok", "data": res_data}

@app.get("/api/teacher/projects/{project_id}")
async def get_teacher_project_detail(project_id: int):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM projects WHERE id = ?", (project_id,))
    row = cursor.fetchone()
    if not row: return {"error": "Project not found"}
    project = dict(row)
    
    cursor.execute("SELECT * FROM project_assessments WHERE project_id = ?", (project_id,))
    pa_row = cursor.fetchone()
    assessment = dict(pa_row) if pa_row else {}
    
    cursor.execute("""
        SELECT id, teacher_name, content, created_at, is_active
        FROM teacher_interventions
        WHERE project_id = ?
        ORDER BY datetime(created_at) DESC, id DESC
        LIMIT 20
    """, (project_id,))
    battle_logs = [dict(r) for r in cursor.fetchall()]
    conn.close()
    return {"project": project, "assessment": assessment, "battle_logs": battle_logs}

@app.post("/api/teacher/interventions")
async def create_teacher_intervention(request: TeacherInterventionRequest):
    if not request.content.strip():
        raise HTTPException(status_code=400, detail="指导建议不能为空")

    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute(
        "INSERT INTO teacher_interventions (project_id, teacher_name, content) VALUES (?, ?, ?)",
        (request.project_id, request.teacher_name.strip(), request.content.strip())
    )
    conn.commit()
    conn.close()
    return {"status": "success", "message": "干预锦囊已下发"}

@app.get("/api/student/notifications")
def get_student_notifications(user_id: str):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    cursor.execute("""
        SELECT DISTINCT
            ti.id,
            ti.project_id,
            p.name AS project_name,
            ti.teacher_name,
            ti.content,
            ti.created_at
        FROM teacher_interventions ti
        JOIN projects p ON p.id = ti.project_id
        LEFT JOIN members m ON m.project_id = p.id
        WHERE ti.is_active = 1
          AND (p.owner_id = ? OR m.student_id = ?)
        ORDER BY datetime(ti.created_at) DESC, ti.id DESC
    """, (user_id, user_id))

    notifications = []
    for row in cursor.fetchall():
        notifications.append({
            "id": row["id"],
            "project_id": row["project_id"],
            "project_name": row["project_name"],
            "teacher_name": row["teacher_name"],
            "content": row["content"],
            "created_at": row["created_at"],
        })

    conn.close()
    return {"items": notifications}

@app.get("/api/teacher/intervention-generator")
async def generate_teaching_plan(teacher_id: Optional[str] = None, teacher_name: Optional[str] = None):
    teacher_identity = (teacher_id or teacher_name or "").strip()
    if teacher_identity:
        return await generate_weekly_intervention_plan(teacher_id=teacher_identity)
    if not teacher_identity:
        return {"plan": "未提供教师身份信息，暂时无法生成教学干预方案。"}

    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    top_mistakes = build_teacher_top_mistakes(cursor, teacher_identity)
    conn.close()
    
    if not top_mistakes:
        return {"plan": "该教师名下暂无项目评估数据，无法生成针对性教学计划。"}
        
    summaries_text = "\\n".join([f"项目[{r['name']}]: {r['audit_summary']}" for r in records if r['audit_summary']])
    if not summaries_text.strip():
        return {"plan": "该教师名下项目的评估数据为空，无法生成计划。"}
    
    try:
        from app.agent.graph import get_llm
        llm = get_llm()
        prompt = f"""你是一个高校创新创业教研主任。以下是【{teacher_name}】老师名下学生项目最近的AI诊断总结：
{summaries_text}

请你根据这些共性错误，生成一份【下周教学干预计划】。要求：
1. 聚焦最高频的1-2个核心谬误。
2. 给出具体的单节课教学安排（例如：理论课讲什么模型，实践课布置什么沙盘任务）。
3. 语气要像专业的教导总结，总字数不要超过300字。
"""
        response = llm.invoke(prompt)
        return {"plan": response.content}
    except Exception as e:
        return {"plan": f"系统生成教学计划失败：{str(e)}"}

@app.get("/api/teacher/weekly-plan")
async def generate_weekly_intervention_plan(teacher_id: Optional[str] = None, teacher_name: Optional[str] = None):
    teacher_identity = (teacher_id or teacher_name or "").strip()
    if not teacher_identity:
        return {"plan": "未提供教师身份信息，暂时无法生成下周干预方案。"}

    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    top_mistakes = build_teacher_top_mistakes(cursor, teacher_identity)
    conn.close()

    if not top_mistakes:
        return {"plan": "当前暂无可用的高频商业逻辑盲区数据，建议先完成项目 AI 审计后再生成下周干预方案。"}

    mistake_summary = "\n".join([
        f"{idx + 1}. {item['name']}，出现 {item['count']} 次；说明：{item['desc']}"
        for idx, item in enumerate(top_mistakes[:3])
    ])

    try:
        from app.agent.graph import get_llm
        llm = get_llm()
        prompt = f"""你是一名高校创新创业课程教研主任。下面是某位指导教师最近在教师端全局项目大盘中统计出的 Top 商业逻辑盲区（H原则冲突）：

{mistake_summary}

请直接生成一份下周干预方案，要求：
1. 只聚焦最高频的1到2个核心错误。
2. 输出必须是教师可直接执行的实操教学计划，而不是泛泛建议。
3. 内容需要覆盖：本周教学目标、课堂讲解重点、课后实操任务、教师观察与验收方式。
4. 语气专业、简洁、有执行感。
5. 总字数控制在300字左右，尽量不要超过320字。
6. 直接输出正文，不要使用 Markdown 列表，不要额外添加标题。
"""
        response = llm.invoke(prompt)
        plan_text = response.content if isinstance(response.content, str) else str(response.content)
        return {"plan": plan_text.strip()}
    except Exception as e:
        return {"plan": f"系统生成教学计划失败：{str(e)}"}

@app.get("/api/teacher/projects/{project_id}/capability_profile")
async def generate_capability_profile(project_id: int):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT owner_id, name FROM projects WHERE id = ?", (project_id,))
    proj = cursor.fetchone()
    if not proj:
        conn.close()
        return {"error": "项目不存在"}
        
    owner_id = proj["owner_id"]
    proj_name = proj["name"]
    
    cursor.execute("""
        SELECT role, content as text
        FROM messages 
        WHERE conversation_id IN (SELECT id FROM conversations WHERE user_id = ?)
        ORDER BY timestamp DESC LIMIT 6
    """, (owner_id,))
    messages_raw = cursor.fetchall()
    conn.close()
    
    if len(messages_raw) < 2:
        return {"error": "对话记录过少，无法生成三轮互动能力画像。请鼓励学生多使用系统。"}
        
    messages_raw.reverse()
    chat_log = ""
    for idx, m in enumerate(messages_raw):
        prefix = "学生" if m["role"] == "user" else "AI导师"
        chat_log += f"[{idx+1}] {prefix}: {m['text']}\\n"
        
    try:
        from app.agent.graph import get_llm
        llm = get_llm()
        prompt = f"""你是一个高级商业教练。以下是负责【{proj_name}】项目的学生与AI导师近期的对抗对话记录：
{chat_log}

请基于这些对话行为，生成格式化的【学生能力画像评估报告】。请严格输出 JSON！格式如下：
{{
    "scores": {{
        "empathy": 1到5的整数(痛点发现能力),
        "ideation": 1到5的整数(方案策划能力),
        "business": 1到5的整数(商业建模能力),
        "execution": 1到5的整数(资源杠杆能力),
        "logic": 1到5的整数(逻辑表达能力)
    }},
    "round_diagnostics": [
        "第一轮表现：...（必须引用至少一句对话日志中的学生原话来作为证据支撑）",
        "第二轮表现：...（必须引用至少一句原话）",
        "综合表现评价：...（总结）"
    ]
}}
"""
        response = llm.invoke(prompt)
        import json
        raw_text = response.content
        if isinstance(raw_text, list):
            raw_text = "".join(str(part) for part in raw_text)
        raw_text = raw_text.strip()
        if "```json" in raw_text:
            raw_text = raw_text.split("```json")[1].split("```")[0].strip()
        result = json.loads(raw_text)
        return {"status": "ok", "profile": result}
    except Exception as e:
        print(f"Failed to generate capability profile: {e}")
        return {"error": f"由于系统或网络异常，生成画像评估失败 ({str(e)})"}

class FinancialAnalysisRequest(BaseModel):
    users: float
    cac: float
    arpu: float
    fixedCost: float
    netProfit: float

@app.post("/api/projects/{project_id}/financial-analysis")
async def project_financial_analysis(project_id: int, request: FinancialAnalysisRequest):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT content FROM projects WHERE id = ?", (project_id,))
    proj = cursor.fetchone()
    conn.close()
    
    project_text = proj["content"] if proj and proj["content"] else "当前项目无核心文档。请要求学生尽快在画布中撰写商业计划。"
    
    try:
        from app.agent.graph import get_llm
        llm = get_llm()
        
        prompt = f"""你现在扮演著名风险投资机构的 AI 高级尽调合伙人 (DeepSeek Financial Pro)。
你正在对一个早期的硬科技或高校双创项目进行残酷的极端压力测试（Financial Due Diligence）。

【项目原始上下文参考】：
{project_text[:2000]}

【创始人提交的当月预测动态财报模型】：
- 初始用户基数: {request.users}人/月
- 综合获客成本 (CAC): {request.cac}元/人
- 单客月度营收 (ARPU): {request.arpu}元
- 硬性固定成本: {request.fixedCost}元/月
- 静态净利润预估: {request.netProfit}元/月

你需要基于以上事实极其残酷地评估。如果净利极其丰厚但商业逻辑薄弱，直接以投资人的身份无情拆穿。
如果参数显示严重亏损，请指出具体止损建议。

请严格返回如下单一纯净的 JSON 格式（不要包含 markdown 代码块标志，只输出大括号里的内容）：
{{
    "risk_assessment": "【危机断言】(不少于60字。指出当期 CAC 与固定成本对现金流的隐蔽消耗)",
    "growth_leverage": "【杠杆测算】(判断 LTV / CAC。指出模型能否长效维持或规模效应的拐点在哪)",
    "next_metric_focus": "【靶向核心】(明确指出创始人下个月必须关注或优化的 1 个关键运营指标编号)"
}}
"""
        response = llm.invoke(prompt)
        raw_text = response.content
        if isinstance(raw_text, list):
            raw_text = "".join(str(part) for part in raw_text)
        raw_text = raw_text.strip()
        if "```json" in raw_text:
            raw_text = raw_text.split("```json")[1].split("```")[0].strip()
        elif raw_text.startswith("```"):
            raw_text = raw_text.strip("```").strip()
            
        import json
        result = json.loads(raw_text)
        return {"status": "success", "advice": result}
    except Exception as e:
        print(f"金融尽调生成失败: {e}")
        return {"status": "error", "message": f"金融沙盘模型崩溃 ({str(e)})"}

# --- Admin Endpoints ---

@app.get("/api/admin/dashboard")
def admin_dashboard():
    """获取管理员全局概览数据"""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    # 1. 项目总数
    cursor.execute("SELECT COUNT(*) as count FROM projects")
    project_count = cursor.fetchone()["count"]
    
    # 2. 附件总数
    cursor.execute("SELECT COUNT(*) as count FROM project_files")
    file_count = cursor.fetchone()["count"]
    
    # 3. 学院分布状况
    cursor.execute("SELECT college, COUNT(*) as count FROM projects GROUP BY college")
    college_dist = {row["college"] or "未知": row["count"] for row in cursor.fetchall()}
    
    # 4. Neo4j 用户总数
    user_count = 0
    try:
        user_result = db.execute_query("MATCH (u:User) RETURN count(u) as count")
        if user_result:
            user_count = user_result[0]["count"]
    except: pass
    
    return {
        "project_count": project_count,
        "file_count": file_count,
        "user_count": user_count,
        "college_distribution": college_dist
    }

@app.get("/api/admin/users")
def admin_get_users():
    """获取 Neo4j 所有实名用户及其学院角色 (管理员视角)"""
    try:
        query = "MATCH (u:User) RETURN u.username as username, u.real_name as real_name, u.role as role, u.college as college"
        result = db.execute_query(query)
        # 显式构造字典列表，解决 JSON 序列化丢失键名导致前端渲染空白的问题
        user_list = []
        for r in result:
            user_list.append({
                "username": r["username"],
                "real_name": r["real_name"],
                "role": r["role"],
                "college": r["college"]
            })
        return user_list
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/admin/identities")
def admin_get_identities():
    """全新的调试接口，确保返回标准的字典列表 (JSON Objects)"""
    try:
        query = "MATCH (u:User) RETURN u.username as username, u.real_name as real_name, u.role as role, u.college as college"
        result = db.execute_query(query)
        return [record.data() for record in result]
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/admin/projects")
def admin_get_projects():
    """获取 SQLite 所有项目及其关联信息"""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    query = """
    SELECT p.*, 
    (SELECT COUNT(*) FROM project_files pf WHERE pf.project_id = p.id) as file_count
    FROM projects p
    """
    cursor.execute(query)
    projects = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return projects

@app.get("/api/admin/sqlite/tables")
def admin_get_sqlite_tables():
    """获取 SQLite 数据库中所有的表名"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'")
    tables = [row[0] for row in cursor.fetchall()]
    conn.close()
    return tables

@app.get("/api/admin/sqlite/data/{table_name}")
def admin_get_sqlite_data(table_name: str):
    """获取指定表的数据预览 (前 100 条)"""
    # 简单的安全过滤，防止 SQL 注入
    if not table_name.isalnum() and "_" not in table_name:
        raise HTTPException(status_code=400, detail="Invalid table name")
        
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    try:
        cursor.execute(f"SELECT * FROM {table_name} LIMIT 100")
        rows = [dict(r) for r in cursor.fetchall()]
        # 获取列定义信息
        cursor.execute(f"PRAGMA table_info({table_name})")
        columns = [{"title": col[1], "key": col[1], "dataIndex": col[1]} for col in cursor.fetchall()]
        return {"columns": columns, "data": rows}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

# --- 用户资料管理 ---
class ProfileUpdate(BaseModel):
    real_name: Optional[str] = None
    college: Optional[str] = None
    major: Optional[str] = None
    grade: Optional[str] = None

@app.get("/api/user/profile")
def get_user_profile(username: str = Query(...)):
    """获取用户详细资料"""
    try:
        query = """
        MATCH (u:User {username: $username})
        RETURN
            u.username as username,
            u.real_name as real_name,
            u.role as role,
            u.college as college,
            u.major as major,
            u.grade as grade
        """
        result = db.execute_query(query, {"username": username})
        if not result:
            raise HTTPException(status_code=404, detail="User not found")
        return record.data() if (record := result[0]) else {}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/user/profile/update")
def update_user_profile(username: str, data: ProfileUpdate):
    """更新用户资料"""
    try:
        # 构建动态更新语句
        updates = []
        params = {"username": username}
        if data.real_name is not None:
            updates.append("u.real_name = $real_name")
            params["real_name"] = data.real_name
        if data.college is not None:
            updates.append("u.college = $college")
            params["college"] = data.college
        if data.major is not None:
            updates.append("u.major = $major")
            params["major"] = data.major
        if data.grade is not None:
            updates.append("u.grade = $grade")
            params["grade"] = data.grade
            
        if not updates:
            return {"message": "No changes"}
            
        query = f"MATCH (u:User {{username: $username}}) SET {', '.join(updates)} RETURN u"
        db.execute_query(query, params)
        return {"message": "Profile updated successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/{full_path:path}", include_in_schema=False)
async def serve_frontend(full_path: str):
    """
    生产部署时由 FastAPI 直接托管 frontend/dist，避免服务器必须安装 Node.js。
    API 路由已在此前声明，这个 catch-all 仅兜底前端静态资源与 SPA 路由。
    """
    if not os.path.isdir(FRONTEND_DIST_DIR):
        raise HTTPException(status_code=404, detail="Frontend dist not found")

    normalized_path = (full_path or "").lstrip("/")
    requested_path = os.path.normpath(os.path.join(FRONTEND_DIST_DIR, normalized_path))
    if normalized_path and requested_path.startswith(FRONTEND_DIST_DIR) and os.path.isfile(requested_path):
        return FileResponse(requested_path)

    index_path = os.path.join(FRONTEND_DIST_DIR, "index.html")
    if os.path.isfile(index_path):
        return FileResponse(index_path)

    raise HTTPException(status_code=404, detail="Frontend index not found")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
