import sys
try:
    from docx import Document
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document

def read_docx(file_path):
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                print(para.text)
        print("\n--- TABLES ---\n")
        for i, table in enumerate(doc.tables):
            print(f"Table {i+1}:")
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip().replace('\n', ' ') for cell in row.cells)
                print(row_text)
            print()
    except Exception as e:
        print(f"Error reading docx: {e}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        read_docx(sys.argv[1])
