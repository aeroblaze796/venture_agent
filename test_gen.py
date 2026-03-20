from docx import Document
import os

def create_test_docx():
    doc = Document()
    doc.add_heading('创新创业项目导入测试', 0)
    doc.add_paragraph('这是一个用于验证 VentureAgent 文件导入功能的测试文档。')
    doc.add_paragraph('项目名称：AI智能农业监测系统')
    doc.add_paragraph('核心技术：计算机视觉、深度学习、物联网。')
    doc.add_paragraph('团队成员：张三（队长）、李四（技术主管）。')
    
    path = os.path.abspath('test_project.docx')
    doc.save(path)
    print(f'Created: {path}')

if __name__ == '__main__':
    create_test_docx()
